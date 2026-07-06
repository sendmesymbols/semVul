from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("reports/SemanticVul_Ablation_Study_Plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9DEE7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, after=6, before=0, line=1.10, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_font(r, size=11, color=INK, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=9.5, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if len(str(value)) < 12 and i != len(row) - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=9.3, color=INK)
    set_table_width(table, widths)
    for row in table.rows:
        row._tr.get_or_add_trPr()
        for cell in row.cells:
            set_cell_margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_labelled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r0 = p.add_run(f"{label}: ")
    set_font(r0, size=10.8, bold=True, color=DARK_BLUE)
    r1 = p.add_run(text)
    set_font(r1, size=10.8, color=INK)
    return p


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("SemanticVul ablation study plan")
    set_font(r, size=9, color=MUTED)
    r2 = p.add_run(" | Page ")
    set_font(r2, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    p._p.append(fld)


def add_masthead(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    add_footer(section)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Plan of Ablation Study")
    set_font(r, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("SemanticVul: Local Explanations, Quality-Aware Fusion, and Imbalance-Aware Evaluation")
    set_font(r, size=13, color=MUTED)

    rows = [
        ("Prepared for", "Supervisor review"),
        ("Purpose", "Response to observation on missing ablation study plan"),
        ("Study focus", "Significance of individual ML pipeline blocks"),
        ("Datasets", "Audited Devign and Reveal splits"),
        ("Date", "July 6, 2026"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for idx, (label, value) in enumerate(rows):
        c0, c1 = table.rows[idx].cells
        for c in (c0, c1):
            set_cell_margins(c, top=50, bottom=50)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_font(r0, size=10.5, bold=True, color=INK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_font(r1, size=10.5, color=INK)
    set_table_width(table, [1900, 7460])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    configure_styles(doc)
    add_masthead(doc)

    add_callout(
        doc,
        "Supervisor observation addressed",
        'Observation: "plan of ablation study is not there." Action to be taken: "The student will perform experiments to report the significance in terms of accuracy/performance measures of individual blocks in the ML Pipeline designed for the problem."',
    )

    add_heading(doc, "1. Objective and Scope", 1)
    add_para(
        doc,
        "This ablation study is designed to make the contribution of each block in the SemanticVul pipeline empirically visible. The study will not only report the final performance of the complete model, but will also show what is gained or lost when each major component is removed, replaced, or simplified under the same experimental protocol.",
    )
    add_para(
        doc,
        "The pipeline is treated as four connected blocks: the code representation channel, the local explanation channel, the label free quality feature channel, and the downstream training and decision strategy. The ablation design therefore combines a bottom-up component ladder with top-down leave-one-out experiments. This allows the report to answer two different but related questions: whether performance improves as components are added, and whether each component remains useful when the full model is already present.",
    )
    add_para(
        doc,
        "All ablations will be conducted on the audited Devign and Reveal splits. The same data partitions, cached representations, evaluation metrics, random seeds, and threshold-selection rules will be used across comparable configurations. This is necessary so that observed performance differences can be attributed to the ablated pipeline block rather than to changes in protocol.",
    )

    add_heading(doc, "2. Research Objectives and Questions", 1)
    add_para(
        doc,
        "The research objectives and research questions are retained word-to-word from the study design. The ablation plan is organized around these objectives so that every experimental comparison has a direct role in the dissertation/report structure.",
    )
    rq_items = [
        (
            "RO 01: Explanations",
            "Design and evaluate a local explanation pipeline using structured JSON output, evidence token grounding, and explicit leakage controls.",
            "To what extent do locally generated, verdict scrubbed, evidence grounded explanations improve explanation faithfulness and downstream vulnerability detection as compared with FuSEVul style free form LLM explanations?",
            "An explanation leakage rate, evidence grounding measures, and explanation quality analysis",
            "A reproducible local explanation generation pipeline and the cached, evidence grounded explanation set with label free quality scores.",
        ),
        (
            "RO 02: Fusion",
            "Develop and evaluate a lightweight gated fusion module that combines cached code embeddings, explanation embeddings, and label free quality features.",
            "How does quality aware adaptive gated fusion compare with static fusion, single modality models, and classical cached feature baselines in vulnerability detection performance and training efficiency?",
            "Controlled ablations, comparison with fusion and single modality baselines, encoder comparison, training efficiency, and predictive performance.",
            "The lightweight SemanticVul gated fusion model over frozen, cached encoders.",
        ),
        (
            "RO 03: Overall Evaluation",
            "Evaluate SemanticVul against FuSEVul and representative baselines under audited, explicitly declared experimental protocols.",
            "Under the audited Devign and Reveal splits, how does SemanticVul compare with FuSEVul and representative baselines in predictive performance, threshold robustness, and low resource training feasibility?",
            "Five seed mean and standard deviation for Accuracy, Precision, Recall, F1, and PR AUC; threshold sensitivity analysis; and data quality findings",
            "A reproducible evaluation framework and comparative benchmark results on Devign and Reveal.",
        ),
        (
            "RO 04: Class Imbalance",
            "Investigate focal loss, capped class weighting, validation based threshold tuning, and multi seed ensembling through controlled ablation experiments.",
            "What are the individual and combined effects of imbalance aware loss functions, validation based threshold tuning, and multi seed ensembling on minority class detection and the precision recall trade off?",
            "Component ablations, minority class Recall, precision recall curves, F1.",
            "An evaluated imbalance handling strategy and a validation based threshold selection procedure.",
        ),
    ]
    for title, objective, question, evidence, output in rq_items:
        add_heading(doc, title, 2)
        add_labelled_paragraph(doc, "Objective", objective)
        add_labelled_paragraph(doc, "Research question", question)
        add_labelled_paragraph(doc, "Evidence to report", evidence)
        add_labelled_paragraph(doc, "Expected output", output)

    add_heading(doc, "3. Components Under Study", 1)
    add_para(
        doc,
        "The ablation study treats the full SemanticVul system as a composition of separable blocks. Each block has a specific expected role in the pipeline and therefore a specific ablation test.",
    )
    comp_rows = [
        ("C", "Code channel", "Cached code embeddings generated from the code encoder.", "Provides the primary source-code representation for vulnerability detection."),
        ("E", "Explanation channel", "Local explanations generated using structured JSON output, verdict scrubbing, and evidence token grounding.", "Adds semantic rationale and local evidence derived from the function."),
        ("Q", "Quality features", "Label free quality features extracted from the explanation and its grounding properties.", "Signals reliability of the explanation to the fusion module."),
        ("G", "Gated fusion", "Quality aware adaptive fusion over cached code embeddings, explanation embeddings, and quality features.", "Learns how much to rely on each channel for a given sample."),
        ("I", "Imbalance strategy", "Focal loss, capped class weighting, validation based threshold tuning, and multi seed ensembling.", "Improves minority class detection and stabilizes the precision recall trade off."),
    ]
    add_table(doc, ["Symbol", "Block", "Operational definition", "Expected contribution"], comp_rows, [800, 1600, 3600, 3360])

    add_heading(doc, "4. Research Questions and Hypotheses", 1)
    add_para(
        doc,
        "The following hypotheses are pre-registered before running the ablations. All configurations will be reported regardless of whether the hypothesis is supported. This prevents selective reporting and keeps the ablation study auditable.",
    )
    hyp_rows = [
        (
            "RQ1",
            "To what extent do locally generated, verdict scrubbed, evidence grounded explanations improve explanation faithfulness and downstream vulnerability detection as compared with FuSEVul style free form LLM explanations?",
            "Structured evidence grounded explanations versus FuSEVul style free form explanations; with and without verdict scrubbing; with and without evidence grounding.",
            "Structured, verdict scrubbed, evidence grounded explanations will reduce leakage, improve grounding, and improve downstream detection compared with free form explanations.",
        ),
        (
            "RQ2",
            "How does quality aware adaptive gated fusion compare with static fusion, single modality models, and classical cached feature baselines in vulnerability detection performance and training efficiency?",
            "Full gated fusion versus static fusion, code-only, explanation-only, code plus quality features, and classical cached feature baselines.",
            "Quality aware adaptive gated fusion will improve F1 and PR AUC over static fusion and single-modality variants while keeping training cost low because encoders are frozen and cached.",
        ),
        (
            "RQ3",
            "Under the audited Devign and Reveal splits, how does SemanticVul compare with FuSEVul and representative baselines in predictive performance, threshold robustness, and low resource training feasibility?",
            "Full SemanticVul model versus FuSEVul and representative baselines under the same audited protocol.",
            "SemanticVul will be competitive with, or stronger than, the comparison baselines and will show more transparent threshold behavior under the audited splits.",
        ),
        (
            "RQ4",
            "What are the individual and combined effects of imbalance aware loss functions, validation based threshold tuning, and multi seed ensembling on minority class detection and the precision recall trade off?",
            "Full imbalance strategy versus removal or isolated addition of focal loss, capped class weighting, threshold tuning, and multi seed ensembling.",
            "Validation based threshold tuning and imbalance aware training will improve minority class Recall and PR AUC; ensembling will mainly improve stability across seeds.",
        ),
    ]
    add_table(doc, ["RQ", "Question", "Comparison", "Hypothesis"], hyp_rows, [700, 3300, 2800, 2560])

    add_heading(doc, "5. Ablation Configurations", 1)
    add_para(
        doc,
        "The core study uses two complementary views. The component ladder starts from the simplest useful detector and adds blocks step by step. The leave-one-out view starts from the complete model and removes one block at a time. Together, these two views report both additive benefit and marginal necessity.",
    )
    ladder_rows = [
        ("L1", "Yes", "No", "No", "No", "Standard", "Code-only baseline."),
        ("L2", "Yes", "Yes", "No", "Static or gated without Q", "Standard", "Measures the contribution of explanation embeddings."),
        ("L3 / Full", "Yes", "Yes", "Yes", "Quality aware adaptive gated fusion", "Full", "Complete SemanticVul configuration."),
        ("-Expl", "Yes", "No", "Yes", "No explanation channel", "Full", "Leave-one-out: removes explanation embeddings."),
        ("-QF", "Yes", "Yes", "No", "Fusion without quality features", "Full", "Leave-one-out: removes quality features. Equivalent to L2 when the same fusion is used."),
        ("Static fusion", "Yes", "Yes", "Yes", "Concatenation or averaging", "Full", "Tests whether adaptive gating improves over naive fusion."),
        ("Explanation only", "No", "Yes", "Optional", "Single modality", "Standard", "Tests whether explanations alone carry predictive signal."),
        ("Classical cached features", "Yes", "No", "Optional", "Classical ML classifier", "Standard", "Lightweight non-neural baseline over cached representations."),
    ]
    add_table(
        doc,
        ["Config", "C", "E", "Q", "Fusion", "Training", "Role"],
        ladder_rows,
        [900, 620, 620, 620, 2100, 1000, 3500],
    )

    add_heading(doc, "6. Explanation-Specific Ablations", 1)
    add_para(
        doc,
        "The explanation ablations are tied to RO 01. They test whether explanation structure and leakage control are doing meaningful work, rather than assuming that any LLM explanation is beneficial.",
    )
    exp_rows = [
        ("Full explanation pipeline", "Structured JSON, verdict scrubbed, evidence grounded.", "Reference configuration for RO 01."),
        ("Free form explanation", "Replace structured JSON explanation with FuSEVul style free form LLM explanation.", "Tests structured output against free form explanation."),
        ("No verdict scrubbing", "Generate or use explanations without verdict scrubbing.", "Measures explanation leakage rate and its effect on downstream detection."),
        ("No evidence grounding", "Remove evidence token grounding from the explanation pipeline.", "Measures faithfulness and grounding contribution."),
        ("No quality scoring", "Keep explanation text but remove label free quality features from fusion.", "Tests whether quality estimates add signal beyond the explanation embedding."),
    ]
    add_table(doc, ["Variant", "Change from full pipeline", "Purpose"], exp_rows, [2200, 4100, 3060])

    add_heading(doc, "7. Imbalance and Threshold Ablations", 1)
    add_para(
        doc,
        "The imbalance ablations are tied to RO 04. They will be run as controlled variants so the report can separate the effect of the loss function, class weighting, threshold selection, and ensembling.",
    )
    imb_rows = [
        ("Base training", "Standard loss, fixed threshold, single seed.", "Reference for the imbalance study."),
        ("+ Focal loss", "Add focal loss only.", "Measures whether hard/minority examples receive useful emphasis."),
        ("+ Capped class weighting", "Add capped class weighting only.", "Measures whether class prior correction improves Recall without excessive false positives."),
        ("+ Threshold tuning", "Use validation based threshold selection only.", "Measures whether decision calibration improves F1 and PR trade off."),
        ("+ Multi seed ensemble", "Use multi seed ensembling only.", "Measures stability and variance reduction."),
        ("Full imbalance strategy", "Combine focal loss, capped class weighting, threshold tuning, and multi seed reporting or ensembling.", "Measures the combined effect of the selected imbalance strategy."),
    ]
    add_table(doc, ["Variant", "Definition", "Purpose"], imb_rows, [2300, 3900, 3160])

    add_heading(doc, "8. Experimental Protocol", 1)
    protocol_points = [
        "Use the audited Devign and Reveal splits for every comparable configuration.",
        "Keep preprocessing, normalization, cached embedding generation, and data filtering fixed across ablations.",
        "Train each main configuration with five random seeds and report mean and standard deviation.",
        "Select thresholds only on the validation split when threshold tuning is enabled.",
        "Report test metrics once per seed using the pre-declared threshold rule.",
        "Record training time, inference time, and number of trainable parameters for training-efficiency analysis.",
        "Report all configurations, including negative or inconclusive results.",
    ]
    for point in protocol_points:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(point)
        set_font(r, size=11, color=INK)

    add_heading(doc, "9. Metrics and Reporting", 1)
    add_para(
        doc,
        "The main predictive metrics will be Accuracy, Precision, Recall, F1, and PR AUC. Because vulnerability datasets are typically imbalanced, minority class Recall, F1, and PR AUC will be treated as central evidence rather than secondary statistics. Accuracy will be reported, but it will not be used alone to justify a model choice.",
    )
    add_para(
        doc,
        "For explanation quality, the report will include explanation leakage rate, evidence grounding measures, and label free explanation quality scores. For robustness, the report will include threshold sensitivity analysis and seed-to-seed variation. For efficiency, the report will include training time and trainable parameter count, with special attention to the benefit of frozen cached encoders.",
    )
    metric_rows = [
        ("Predictive performance", "Accuracy, Precision, Recall, F1, PR AUC", "All ROs, especially RO 02 and RO 03"),
        ("Minority class behavior", "Minority class Recall, PR curves, F1", "RO 04"),
        ("Explanation faithfulness", "Evidence grounding measures and quality analysis", "RO 01"),
        ("Leakage control", "Explanation leakage rate", "RO 01"),
        ("Robustness", "Five seed mean and standard deviation; threshold sensitivity", "RO 03 and RO 04"),
        ("Efficiency", "Training time, inference time, trainable parameters", "RO 02 and RO 03"),
    ]
    add_table(doc, ["Evidence type", "Metric or artifact", "Used for"], metric_rows, [2500, 3600, 3260])

    add_heading(doc, "10. Interpretation Plan", 1)
    add_para(
        doc,
        "A component will be considered useful if its inclusion produces a consistent improvement in F1, PR AUC, or minority class Recall without causing an unacceptable increase in leakage, training cost, or threshold instability. The final interpretation will not rely on a single metric. For example, a configuration that improves Accuracy while reducing minority class Recall will be treated cautiously because it may be exploiting the majority class distribution.",
    )
    add_para(
        doc,
        "The report will present both absolute scores and deltas from the relevant reference configuration. For the ladder, the key deltas are L2 minus L1 and Full minus L2. For leave-one-out, the key values are Full minus each ablated variant. These deltas directly show the significance of each individual block in the pipeline.",
    )

    add_callout(
        doc,
        "Expected final deliverable",
        "The final ablation section will contain one ladder table, one leave-one-out table, one explanation ablation table, one imbalance ablation table, and a short interpretation paragraph for each dataset. This directly satisfies the requested action: reporting the significance of individual ML pipeline blocks in terms of accuracy and performance measures.",
    )

    doc.core_properties.title = "Plan of Ablation Study for SemanticVul"
    doc.core_properties.subject = "Supervisor review document"
    doc.core_properties.author = "SemanticVul student researcher"
    doc.core_properties.keywords = "SemanticVul, ablation study, vulnerability detection, explanations, fusion"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build().resolve())
